import io
import requests
import os
import google.generativeai as genai  
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app) 

SAPLING_API_KEY = "H5TADY2DJHOAM9JZI3WPD21AIBOXMNF4"

OLLAMA_URL = "http://localhost:11434/api/generate"
OLLAMA_MODEL = "deepseek-r1:1.5b"


GEMINI_API_KEY = "AIzaSyBhCIgYVRxjrwYakuNfFSvSr8atVcvhghg" 
genai.configure(api_key=GEMINI_API_KEY)


def generate_text_ollama(prompt, temperature=0.7):
    """Generates the initial text using local Ollama."""
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature,
            "top_k": 40,
            "top_p": 0.9
        }
    }
    try:
        response = requests.post(OLLAMA_URL, json=payload)
        response.raise_for_status()
        text = response.json().get("response", "")
        
        if "</think>" in text:
            text = text.split("</think>")[-1].strip()
        return text
    except requests.exceptions.ConnectionError:
        print("Error: Cannot connect to Ollama. Is 'ollama serve' running?")
        return None
    except Exception as e:
        print(f"Ollama Error: {e}")
        return None

def humanize_with_gemini(text):
    """Uses Gemini to rewrite text to be more human-like."""
    try:
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        humanize_prompt = """
### CORE INSTRUCTIONS
You are a casual, relatable human writer. Your goal is to rewrite the provided text so it is indistinguishable from human writing. You must abandon all "AI-style" tendencies.

### STYLE & TONE GUIDELINES
1.  Burstiness & Flow: Do not write distinct, uniform paragraphs. Mix very short, punchy sentences with longer, flowing ones. Mimic the natural rhythm of human thought.
2.  Imperfection: Humans are not perfectly concise. Use occasional colloquialisms, idioms, or conversational transitions (e.g., "Honestly," "Look," "Here's the thing").
3.  Subjectivity: AI is neutral; humans are opinionated. Unless strictly formal, lean into a subjective perspective. Use "I," "we," and active voice.
4.  Simplicity: Avoid complex vocabulary where simple words suffice. (e.g., use "use" instead of "utilize", "help" instead of "facilitate").

### NEGATIVE CONSTRAINTS (STRICTLY FORBIDDEN)
1.  No "AI Words": diverse, realm, landscape, tapestry, delve, underscore, crucial, pivotal, intricate, unleash, game-changer, dynamic, foster, merely, arguably.
2.  No Structural Fluff: deeply avoid transition words like "Furthermore," "Moreover," "In conclusion," "Additionally," "Firstly/Secondly."
3.  No Em Dashes: Do not use the em-dash sign (—). Use commas, periods, or parentheses instead.
4.  No Markdown formatting: Do not use bolding (**text**), italics, or headers unless explicitly asked. Write as raw text.
5.  No Lists: Avoid bullet points. Humans usually write in paragraphs, not lists.

### FORMATTING
-   Write in a continuous flow.
-   Keep the output raw and conversational.
-   Do not preach or moralize at the end.
"""
        
        full_prompt = f"{humanize_prompt}\n\nHere is the text to rewrite:\n{text}"
        
        response = model.generate_content(full_prompt)
        return response.text
    except Exception as e:
        print(f"Gemini API Error: {e}")
        return None

def check_ai_score(text):
    url = "https://api.sapling.ai/api/v1/aidetect"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {SAPLING_API_KEY}"
    }
    try:
        response = requests.post(url, json={"text": text}, headers=headers)
        response.raise_for_status()
        return response.json().get("score", 0)
    except Exception as e:
        print(f"Sapling Error: {e}")
        return 0.0

def add_page_borders(section):
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '24')
        border.set(qn('w:color'), '000000')
        pgBorders.append(border)
    sectPr.append(pgBorders)


@app.route('/')
def home():
    return send_file('index.html')

@app.route('/generate', methods=['POST'])
def api_generate():
    """Initial generation using Ollama"""
    data = request.json
    prompt = data.get('prompt')
    response = generate_text_ollama(prompt, temperature=0.7)
    
    if response is None:
        return jsonify({"error": "Could not connect to Ollama. Please make sure Ollama is running."}), 503
        
    return jsonify({"response": response})

@app.route('/check_score', methods=['POST'])
def api_check_score():
    data = request.json
    text = data.get('text')
    score = check_ai_score(text)
    return jsonify({"score": score})

@app.route('/humanize', methods=['POST'])
def api_humanize():
    """Refining generation using Gemini"""
    data = request.json
    text = data.get('text') 
    
    if not text:
        return jsonify({"error": "No text provided to humanize"}), 400

    response = humanize_with_gemini(text)
    
    if response is None:
         return jsonify({"error": "Gemini API failed or is unreachable."}), 503

    return jsonify({"response": response})

@app.route('/download', methods=['POST'])
def api_download():
    data = request.json
    content = data.get('content')
    header_text = data.get('header')
    footer_text = data.get('footer')

    doc = Document()
    section = doc.sections[0]
    
    header = section.header
    header.paragraphs[0].text = header_text
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    add_page_borders(section)
    
    doc.add_heading('Generated Submission', 0)
    doc.add_paragraph(content)
    
    footer = section.footer
    footer.paragraphs[0].text = footer_text
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name="submission.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == '__main__':
    print("Starting Flask Server on port 5000...")
    app.run(debug=True, port=5000)